"""
DOSSIER — Text Extraction Pipeline
Handles: PDF, plain text, HTML, images (OCR), DOCX, XLSX, CSV, JSON, EML, RTF, ZIP.
"""

import csv
import email
import hashlib
import io
import json
import os
import subprocess
import tempfile
import zipfile
from pathlib import Path


def file_hash(filepath: str) -> str:
    """SHA-256 hash for dedup."""
    h = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


# All supported extensions
SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".txt",
    ".md",
    ".log",
    ".html",
    ".htm",
    ".png",
    ".jpg",
    ".jpeg",
    ".tiff",
    ".tif",
    ".bmp",
    ".gif",
    ".webp",
    ".docx",
    ".xlsx",
    ".xls",
    ".csv",
    ".tsv",
    ".json",
    ".jsonl",
    ".eml",
    ".msg",
    ".rtf",
    ".odt",
    ".xml",
    ".zip",
}


def extract_text(filepath: str) -> dict:
    """
    Extract text from a file. Returns dict with:
      - text: extracted text content
      - pages: page count (if applicable)
      - method: extraction method used
    """
    path = Path(filepath)
    suffix = path.suffix.lower()

    extractors = {
        ".pdf": _extract_pdf,
        ".txt": _extract_text,
        ".md": _extract_text,
        ".log": _extract_text,
        ".html": _extract_html,
        ".htm": _extract_html,
        ".xml": _extract_html,
        ".png": _extract_image_ocr,
        ".jpg": _extract_image_ocr,
        ".jpeg": _extract_image_ocr,
        ".tiff": _extract_image_ocr,
        ".tif": _extract_image_ocr,
        ".bmp": _extract_image_ocr,
        ".gif": _extract_image_ocr,
        ".webp": _extract_image_ocr,
        ".docx": _extract_docx,
        ".xlsx": _extract_xlsx,
        ".xls": _extract_xlsx,
        ".csv": _extract_csv,
        ".tsv": _extract_tsv,
        ".json": _extract_json,
        ".jsonl": _extract_jsonl,
        ".eml": _extract_eml,
        ".msg": _extract_eml,
        ".rtf": _extract_rtf,
        ".odt": _extract_odt,
    }

    extractor = extractors.get(suffix)
    if not extractor:
        return {"text": "", "pages": 0, "method": "unsupported"}

    return extractor(filepath)


def _extract_pdf(filepath: str) -> dict:
    """Extract from PDF — tries native text first, falls back to OCR."""
    import pdfplumber

    text_parts = []
    pages = 0

    try:
        with pdfplumber.open(filepath) as pdf:
            pages = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        print(f"[EXTRACT] pdfplumber failed on {filepath}: {e}")

    full_text = "\n\n".join(text_parts).strip()

    # If we got very little text, the PDF is probably scanned — try OCR
    if len(full_text) < 100 and pages > 0:
        print(f"[EXTRACT] Low text yield ({len(full_text)} chars), attempting OCR...")
        ocr_result = _ocr_pdf(filepath)
        if len(ocr_result) > len(full_text):
            return {"text": ocr_result, "pages": pages, "method": "pdf_ocr"}

    return {"text": full_text, "pages": pages, "method": "pdf_native"}


def _ocr_pdf(filepath: str) -> str:
    """OCR a PDF by converting pages to images then running tesseract."""
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            subprocess.run(
                ["pdftoppm", "-png", "-r", "300", filepath, os.path.join(tmpdir, "page")],
                capture_output=True,
                timeout=120,
            )

            text_parts = []
            for img_file in sorted(Path(tmpdir).glob("*.png")):
                result = subprocess.run(
                    ["tesseract", str(img_file), "stdout", "--psm", "3"],
                    capture_output=True,
                    text=True,
                    timeout=60,
                )
                if result.stdout.strip():
                    text_parts.append(result.stdout.strip())

            return "\n\n".join(text_parts)
    except Exception as e:
        print(f"[OCR] Failed: {e}")
        return ""


def _extract_text(filepath: str) -> dict:
    """Plain text / markdown / log files."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        pages = max(1, len(text) // 3000)
        return {"text": text, "pages": pages, "method": "plaintext"}
    except Exception as e:
        print(f"[EXTRACT] Text read failed: {e}")
        return {"text": "", "pages": 0, "method": "plaintext_error"}


def _extract_html(filepath: str) -> dict:
    """Strip HTML/XML tags, extract text content."""
    import re

    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            html = f.read()

        html = re.sub(r"<(script|style)[^>]*>.*?</\1>", "", html, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", html)
        text = re.sub(r"\s+", " ", text).strip()

        pages = max(1, len(text) // 3000)
        return {"text": text, "pages": pages, "method": "html"}
    except Exception as e:
        print(f"[EXTRACT] HTML parse failed: {e}")
        return {"text": "", "pages": 0, "method": "html_error"}


def _extract_image_ocr(filepath: str) -> dict:
    """OCR an image file directly."""
    try:
        result = subprocess.run(
            ["tesseract", filepath, "stdout", "--psm", "3"],
            capture_output=True,
            text=True,
            timeout=60,
        )
        text = result.stdout.strip()
        return {"text": text, "pages": 1, "method": "image_ocr"}
    except Exception as e:
        print(f"[OCR] Image OCR failed: {e}")
        return {"text": "", "pages": 1, "method": "image_ocr_error"}


def _extract_docx(filepath: str) -> dict:
    """Extract text from Microsoft Word .docx files."""
    try:
        from docx import Document

        doc = Document(filepath)
        parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        text = "\n".join(parts)
        pages = max(1, len(text) // 3000)
        return {"text": text, "pages": pages, "method": "docx"}
    except Exception as e:
        print(f"[EXTRACT] DOCX failed: {e}")
        return {"text": "", "pages": 0, "method": "docx_error"}


def _extract_xlsx(filepath: str) -> dict:
    """Extract text from Excel .xlsx/.xls files."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(filepath, read_only=True, data_only=True)
        parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                row_vals = [str(cell) if cell is not None else "" for cell in row]
                row_text = " | ".join(v for v in row_vals if v)
                if row_text.strip():
                    parts.append(row_text)

        wb.close()
        text = "\n".join(parts)
        pages = max(1, len(text) // 3000)
        return {"text": text, "pages": pages, "method": "xlsx"}
    except Exception as e:
        print(f"[EXTRACT] XLSX failed: {e}")
        return {"text": "", "pages": 0, "method": "xlsx_error"}


def _extract_csv(filepath: str) -> dict:
    """Extract text from CSV files."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            parts = []
            for i, row in enumerate(reader):
                if i > 50000:  # Safety limit
                    parts.append(f"... truncated at {i} rows ...")
                    break
                row_text = " | ".join(cell.strip() for cell in row if cell.strip())
                if row_text:
                    parts.append(row_text)

        text = "\n".join(parts)
        pages = max(1, len(text) // 3000)
        return {"text": text, "pages": pages, "method": "csv"}
    except Exception as e:
        print(f"[EXTRACT] CSV failed: {e}")
        return {"text": "", "pages": 0, "method": "csv_error"}


def _extract_tsv(filepath: str) -> dict:
    """Extract text from TSV files."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f, delimiter="\t")
            parts = []
            for i, row in enumerate(reader):
                if i > 50000:
                    parts.append(f"... truncated at {i} rows ...")
                    break
                row_text = " | ".join(cell.strip() for cell in row if cell.strip())
                if row_text:
                    parts.append(row_text)

        text = "\n".join(parts)
        pages = max(1, len(text) // 3000)
        return {"text": text, "pages": pages, "method": "tsv"}
    except Exception as e:
        print(f"[EXTRACT] TSV failed: {e}")
        return {"text": "", "pages": 0, "method": "tsv_error"}


def _extract_json(filepath: str) -> dict:
    """Extract text from JSON files — flattens structure to readable text."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            data = json.load(f)

        text = _flatten_json(data)
        pages = max(1, len(text) // 3000)
        return {"text": text, "pages": pages, "method": "json"}
    except Exception as e:
        print(f"[EXTRACT] JSON failed: {e}")
        return {"text": "", "pages": 0, "method": "json_error"}


def _extract_jsonl(filepath: str) -> dict:
    """Extract text from JSON Lines files."""
    try:
        parts = []
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            for i, line in enumerate(f):
                if i > 50000:
                    parts.append(f"... truncated at {i} lines ...")
                    break
                line = line.strip()
                if not line:
                    continue
                try:
                    data = json.loads(line)
                    parts.append(_flatten_json(data))
                except json.JSONDecodeError:
                    parts.append(line)

        text = "\n---\n".join(parts)
        pages = max(1, len(text) // 3000)
        return {"text": text, "pages": pages, "method": "jsonl"}
    except Exception as e:
        print(f"[EXTRACT] JSONL failed: {e}")
        return {"text": "", "pages": 0, "method": "jsonl_error"}


def _flatten_json(data, prefix: str = "", max_depth: int = 10) -> str:
    """Recursively flatten JSON to readable key: value lines."""
    if max_depth <= 0:
        return str(data)[:200]

    lines = []
    if isinstance(data, dict):
        for key, value in data.items():
            full_key = f"{prefix}.{key}" if prefix else key
            if isinstance(value, (dict, list)):
                lines.append(_flatten_json(value, full_key, max_depth - 1))
            else:
                lines.append(f"{full_key}: {value}")
    elif isinstance(data, list):
        for i, item in enumerate(data[:500]):  # Safety limit
            lines.append(_flatten_json(item, f"{prefix}[{i}]", max_depth - 1))
    else:
        lines.append(f"{prefix}: {data}" if prefix else str(data))

    return "\n".join(lines)


def _extract_eml(filepath: str) -> dict:
    """Extract text from email .eml files."""
    try:
        with open(filepath, "rb") as f:
            msg = email.message_from_binary_file(f)

        parts = []

        # Headers
        for header in ["From", "To", "Cc", "Bcc", "Subject", "Date", "Message-ID"]:
            value = msg.get(header)
            if value:
                parts.append(f"{header}: {value}")

        parts.append("")  # Blank line separator

        # Body
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        charset = part.get_content_charset() or "utf-8"
                        parts.append(payload.decode(charset, errors="replace"))
                elif content_type == "text/html":
                    payload = part.get_payload(decode=True)
                    if payload:
                        charset = part.get_content_charset() or "utf-8"
                        html_text = payload.decode(charset, errors="replace")
                        # Strip HTML tags
                        import re

                        clean = re.sub(r"<[^>]+>", " ", html_text)
                        clean = re.sub(r"\s+", " ", clean).strip()
                        parts.append(clean)
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                charset = msg.get_content_charset() or "utf-8"
                parts.append(payload.decode(charset, errors="replace"))

        text = "\n".join(parts)
        pages = max(1, len(text) // 3000)
        return {"text": text, "pages": pages, "method": "eml"}
    except Exception as e:
        print(f"[EXTRACT] EML failed: {e}")
        return {"text": "", "pages": 0, "method": "eml_error"}


def _extract_rtf(filepath: str) -> dict:
    """Extract text from RTF files."""
    try:
        from striprtf.striprtf import rtf_to_text

        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            rtf_content = f.read()

        text = rtf_to_text(rtf_content)
        pages = max(1, len(text) // 3000)
        return {"text": text, "pages": pages, "method": "rtf"}
    except Exception as e:
        print(f"[EXTRACT] RTF failed: {e}")
        return {"text": "", "pages": 0, "method": "rtf_error"}


def _extract_odt(filepath: str) -> dict:
    """Extract text from OpenDocument .odt files (they're ZIP-based XML)."""
    try:
        import re

        with zipfile.ZipFile(filepath, "r") as z:
            if "content.xml" not in z.namelist():
                return {"text": "", "pages": 0, "method": "odt_error"}

            content = z.read("content.xml").decode("utf-8", errors="replace")

        # Strip XML tags
        text = re.sub(r"<[^>]+>", " ", content)
        text = re.sub(r"\s+", " ", text).strip()
        pages = max(1, len(text) // 3000)
        return {"text": text, "pages": pages, "method": "odt"}
    except Exception as e:
        print(f"[EXTRACT] ODT failed: {e}")
        return {"text": "", "pages": 0, "method": "odt_error"}


def extract_zip(filepath: str, dest_dir: str) -> list[str]:
    """
    Extract a ZIP file to dest_dir. Returns list of extracted file paths.
    Handles nested ZIPs by extracting recursively.
    """
    extracted = []
    try:
        with zipfile.ZipFile(filepath, "r") as z:
            # Security: check for zip bombs and path traversal
            total_size = sum(info.file_size for info in z.infolist())
            if total_size > 5 * 1024 * 1024 * 1024:  # 5GB limit
                print(f"[ZIP] Skipping {filepath}: extracted size would exceed 5GB")
                return []

            for info in z.infolist():
                # Skip directories and hidden files
                if info.is_dir():
                    continue
                if info.filename.startswith("__MACOSX") or info.filename.startswith("."):
                    continue

                # Prevent path traversal
                safe_name = os.path.basename(info.filename)
                if not safe_name:
                    continue

                # Preserve subdirectory structure
                rel_dir = os.path.dirname(info.filename)
                target_dir = os.path.join(dest_dir, rel_dir) if rel_dir else dest_dir
                os.makedirs(target_dir, exist_ok=True)

                target_path = os.path.join(target_dir, safe_name)

                # Extract the file
                with z.open(info) as src, open(target_path, "wb") as dst:
                    dst.write(src.read())

                # If it's a nested ZIP, extract it too
                if safe_name.lower().endswith(".zip"):
                    nested_dir = os.path.join(target_dir, Path(safe_name).stem)
                    os.makedirs(nested_dir, exist_ok=True)
                    nested_files = extract_zip(target_path, nested_dir)
                    extracted.extend(nested_files)
                else:
                    extracted.append(target_path)

    except zipfile.BadZipFile:
        print(f"[ZIP] Bad zip file: {filepath}")
    except Exception as e:
        print(f"[ZIP] Failed to extract {filepath}: {e}")

    return extracted
